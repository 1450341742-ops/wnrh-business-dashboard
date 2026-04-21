# -*- coding: utf-8 -*-
"""万宁睿和经营分析驾驶舱 Enterprise V2.0
Streamlit + SQLite。首次登录：admin / admin123
"""
from __future__ import annotations

import io, os, re, sqlite3, hashlib, hmac, secrets, zipfile
from datetime import date, datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
import streamlit as st
from docx import Document

APP_TITLE = "万宁睿和经营分析驾驶舱"
DB_PATH = Path("data/business_dashboard.db")
DB_PATH.parent.mkdir(exist_ok=True)

TABLES = {
    "projects": {
        "label": "项目经营明细",
        "cols": ["project_id","project_name","customer","project_type","phase","disease_area","contract_amount","recognized_revenue","project_visits","start_date","end_date","pm","bd","status"],
        "required": ["project_name"],
    },
    "payments": {
        "label": "合同与回款台账",
        "cols": ["contract_id","project_id","project_name","customer","contract_amount","invoice_amount","received_amount","due_date","received_date","bd","owner","remark"],
        "required": ["project_name"],
    },
    "people": {
        "label": "人员成本设置",
        "cols": ["name","role","employment_type","monthly_salary","social_security","bonus_alloc","management_alloc","monthly_total_cost","daily_cost","standard_monthly_visits","can_lead","level","city"],
        "required": ["name"],
    },
    "schedules": {
        "label": "排班与工时明细",
        "cols": ["project_id","project_name","site_name","visit_date","auditor","audit_role","work_days","travel_city","is_parttime","visit_equivalent","remark"],
        "required": ["project_name","auditor"],
    },
    "travel": {
        "label": "差旅原始数据",
        "cols": ["travel_date","person","project_id","project_name","customer","city","expense_type","amount","pay_method","invoice_status","source","remark"],
        "required": ["amount"],
    },
    "params": {
        "label": "基础参数设置",
        "cols": ["key","value","remark"],
        "required": ["key"],
    },
}

ALIASES = {
    "project_id": ["项目编号","项目ID","project_id","编号"],
    "project_name": ["项目名称","项目","project_name","试验项目"],
    "customer": ["客户名称","客户","申办方","委托方","customer"],
    "project_type": ["项目类型","业务类型","类型"],
    "phase": ["分期","临床分期","phase"],
    "disease_area": ["疾病领域","适应症","disease_area"],
    "contract_amount": ["合同金额","合同额","contract_amount","收入"],
    "recognized_revenue": ["确认收入","已确认收入","recognized_revenue"],
    "project_visits": ["项目院次数","院次","院次数","project_visits"],
    "start_date": ["启动日期","开始日期","项目启动日期","start_date"],
    "end_date": ["完成日期","结束日期","项目完成日期","end_date"],
    "pm": ["项目负责人","PM","pm"],
    "bd": ["商务负责人","BD","bd"],
    "status": ["项目状态","状态","status"],
    "contract_id": ["合同编号","合同ID","contract_id"],
    "invoice_amount": ["开票金额","已开票金额","invoice_amount"],
    "received_amount": ["已回款金额","到账金额","回款金额","received_amount"],
    "due_date": ["应回款日期","预计回款日期","due_date"],
    "received_date": ["实际回款日期","到账日期","received_date"],
    "owner": ["催收责任人","负责人","owner"],
    "name": ["姓名","人员","稽查员","name"],
    "role": ["岗位","角色","role"],
    "employment_type": ["人员类型","是否全职","employment_type"],
    "monthly_salary": ["月工资","工资","monthly_salary"],
    "social_security": ["社保公积金","社保","social_security"],
    "bonus_alloc": ["奖金分摊","bonus_alloc"],
    "management_alloc": ["管理分摊","管理成本分摊","management_alloc"],
    "monthly_total_cost": ["月综合成本","monthly_total_cost"],
    "daily_cost": ["日成本","daily_cost"],
    "standard_monthly_visits": ["标准月产能院次","月标准院次","standard_monthly_visits"],
    "can_lead": ["是否可带组","可带组","can_lead"],
    "level": ["等级","职级","level"],
    "city": ["城市","所在城市","city"],
    "site_name": ["中心名称","医院","site_name"],
    "visit_date": ["院次日期","日期","visit_date"],
    "auditor": ["稽查员","人员","auditor"],
    "audit_role": ["稽查角色","项目角色","audit_role"],
    "work_days": ["工作天数","人天","work_days"],
    "travel_city": ["出差城市","城市","travel_city"],
    "is_parttime": ["是否兼职","兼职","is_parttime"],
    "visit_equivalent": ["院次折算","折算院次","visit_equivalent"],
    "travel_date": ["费用日期","差旅日期","日期","travel_date"],
    "person": ["报销人","出差人","人员","person"],
    "expense_type": ["费用类型","科目","expense_type"],
    "amount": ["金额","费用金额","amount"],
    "pay_method": ["支付方式","pay_method"],
    "invoice_status": ["发票状态","invoice_status"],
    "source": ["数据来源","来源","source"],
    "remark": ["备注","说明","remark"],
    "key": ["参数","key"],
    "value": ["值","value"],
}

ROLES = {
    "超级管理员": ["all"],
    "老板": ["view","export","report"],
    "管理层": ["view","edit","import","export","report"],
    "财务": ["view","edit","import","export","report"],
    "项目经理": ["view","edit","import","export"],
    "商务": ["view","edit","export"],
    "只读": ["view"],
}


def conn():
    return sqlite3.connect(DB_PATH)


def q(sql: str, params=()):
    with conn() as c:
        return pd.read_sql_query(sql, c, params=params)


def exec_sql(sql: str, params=()):
    with conn() as c:
        c.execute(sql, params)
        c.commit()


def hash_pw(password: str, salt: str | None = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120000).hex()
    return f"{salt}${digest}"


def verify_pw(password: str, stored: str) -> bool:
    try:
        salt, digest = stored.split("$", 1)
        return hmac.compare_digest(hash_pw(password, salt).split("$",1)[1], digest)
    except Exception:
        return False


def init_db():
    with conn() as c:
        c.execute("CREATE TABLE IF NOT EXISTS users(username TEXT PRIMARY KEY, password_hash TEXT, role TEXT, active INTEGER DEFAULT 1)")
        c.execute("CREATE TABLE IF NOT EXISTS audit_logs(ts TEXT, username TEXT, action TEXT, detail TEXT)")
        for t, meta in TABLES.items():
            cols = ", ".join([f'"{col}" TEXT' for col in meta["cols"]])
            c.execute(f'CREATE TABLE IF NOT EXISTS {t}(id INTEGER PRIMARY KEY AUTOINCREMENT, {cols})')
        cnt = c.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        if cnt == 0:
            c.execute("INSERT INTO users(username,password_hash,role,active) VALUES(?,?,?,1)", ("admin", hash_pw("admin123"), "超级管理员"))
            c.execute("INSERT INTO params(key,value,remark) VALUES(?,?,?)", ("target_margin_rate", "0.40", "目标毛利率"))
            c.execute("INSERT INTO params(key,value,remark) VALUES(?,?,?)", ("tax_rate", "0.06", "税率"))
        c.commit()


def log(action: str, detail: str = ""):
    exec_sql("INSERT INTO audit_logs VALUES(?,?,?,?)", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), st.session_state.get("user",""), action, detail[:500]))


def has_perm(p: str) -> bool:
    role = st.session_state.get("role", "只读")
    perms = ROLES.get(role, [])
    return "all" in perms or p in perms


def read_table(name: str) -> pd.DataFrame:
    return q(f"SELECT * FROM {name}").drop(columns=["id"], errors="ignore")


def write_table(name: str, df: pd.DataFrame, mode="replace"):
    df = df.copy()
    for col in TABLES[name]["cols"]:
        if col not in df.columns:
            df[col] = ""
    df = df[TABLES[name]["cols"]].fillna("")
    with conn() as c:
        if mode == "replace":
            c.execute(f"DELETE FROM {name}")
        df.to_sql(name, c, if_exists="append", index=False)
    log("写入数据", f"{TABLES[name]['label']} {len(df)}行")


def num(s) -> float:
    return pd.to_numeric(s, errors="coerce").fillna(0) if isinstance(s, pd.Series) else float(pd.to_numeric(pd.Series([s]), errors="coerce").fillna(0).iloc[0])


def dt(s):
    return pd.to_datetime(s, errors="coerce")


def normalize_columns(df: pd.DataFrame, table: str) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    rename = {}
    for target in TABLES[table]["cols"]:
        for a in ALIASES.get(target, [target]):
            if a in df.columns:
                rename[a] = target
                break
    df = df.rename(columns=rename)
    for c in TABLES[table]["cols"]:
        if c not in df.columns:
            df[c] = ""
    return df[TABLES[table]["cols"]]


def load_upload(file) -> pd.DataFrame:
    name = file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(file)
    return pd.read_excel(file)


def infer_expense_type(text: str) -> str:
    t = str(text)
    rules = [("住宿","酒店|华住|住宿|宾馆"),("机票/高铁","机票|航司|高铁|火车|铁路|动车"),("市内交通","滴滴|出租|打车|网约车|地铁"),("餐补/餐费","餐|饭|补贴"),("快递/打印","快递|打印|复印")]
    for label, pat in rules:
        if re.search(pat, t): return label
    return "其他"


def auto_recognize_travel(df: pd.DataFrame) -> pd.DataFrame:
    raw = df.copy()
    text_cols = [c for c in raw.columns if raw[c].dtype == object]
    def find_col(keys):
        for c in raw.columns:
            if any(k in str(c) for k in keys): return c
        return None
    amount_col = find_col(["金额","费用","支出","借方"])
    person_col = find_col(["姓名","报销人","出差人","人员"])
    date_col = find_col(["日期","时间"])
    project_col = find_col(["项目"])
    city_col = find_col(["城市","地点"])
    rows = []
    for _, r in raw.iterrows():
        combined = " ".join([str(r.get(c,"")) for c in raw.columns])
        rows.append({
            "travel_date": r.get(date_col, "") if date_col else "",
            "person": r.get(person_col, "") if person_col else "",
            "project_id": "",
            "project_name": r.get(project_col, "") if project_col else "",
            "customer": "",
            "city": r.get(city_col, "") if city_col else "",
            "expense_type": infer_expense_type(combined),
            "amount": r.get(amount_col, 0) if amount_col else 0,
            "pay_method": "",
            "invoice_status": "",
            "source": "自动识别",
            "remark": combined[:120],
        })
    return pd.DataFrame(rows)[TABLES["travel"]["cols"]]


def calc_all():
    projects = read_table("projects")
    payments = read_table("payments")
    people = read_table("people")
    schedules = read_table("schedules")
    travel = read_table("travel")

    for df in [projects, payments, people, schedules, travel]:
        if df.empty: continue

    if projects.empty:
        projects = pd.DataFrame(columns=TABLES["projects"]["cols"])
    projects["revenue"] = num(projects.get("recognized_revenue", 0))
    projects.loc[projects["revenue"].eq(0), "revenue"] = num(projects.get("contract_amount", 0))
    projects["visits"] = num(projects.get("project_visits", 0))

    if not people.empty:
        people["monthly_total_cost_calc"] = num(people.get("monthly_total_cost",0))
        fallback = num(people.get("monthly_salary",0)) + num(people.get("social_security",0)) + num(people.get("bonus_alloc",0)) + num(people.get("management_alloc",0))
        people.loc[people["monthly_total_cost_calc"].eq(0), "monthly_total_cost_calc"] = fallback
        people["daily_cost_calc"] = num(people.get("daily_cost",0))
        people.loc[people["daily_cost_calc"].eq(0), "daily_cost_calc"] = people["monthly_total_cost_calc"] / 21.75
        people_cost = people[["name","daily_cost_calc","monthly_total_cost_calc","standard_monthly_visits"]]
    else:
        people_cost = pd.DataFrame(columns=["name","daily_cost_calc","monthly_total_cost_calc","standard_monthly_visits"])

    if not schedules.empty:
        schedules["work_days_n"] = num(schedules.get("work_days",0))
        schedules["visit_equivalent_n"] = num(schedules.get("visit_equivalent",0))
        schedules.loc[schedules["visit_equivalent_n"].eq(0), "visit_equivalent_n"] = 1
        sch = schedules.merge(people_cost, left_on="auditor", right_on="name", how="left")
        sch["daily_cost_calc"] = num(sch.get("daily_cost_calc",0))
        sch["labor_cost"] = sch["work_days_n"] * sch["daily_cost_calc"]
        labor = sch.groupby("project_name", as_index=False)["labor_cost"].sum()
        capacity = sch.groupby("auditor", as_index=False).agg(visits=("visit_equivalent_n","sum"), travel_days=("work_days_n","sum"), labor_cost=("labor_cost","sum"))
    else:
        sch = pd.DataFrame(); labor = pd.DataFrame(columns=["project_name","labor_cost"]); capacity = pd.DataFrame(columns=["auditor","visits","travel_days","labor_cost"])

    if not travel.empty:
        travel["amount_n"] = num(travel.get("amount",0))
        travel_sum = travel.groupby("project_name", as_index=False)["amount_n"].sum().rename(columns={"amount_n":"travel_cost"})
    else:
        travel_sum = pd.DataFrame(columns=["project_name","travel_cost"])

    if not payments.empty:
        payments["received_n"] = num(payments.get("received_amount",0))
        payments["contract_n"] = num(payments.get("contract_amount",0))
        pay_sum = payments.groupby("project_name", as_index=False).agg(received=("received_n","sum"), contract=("contract_n","sum"))
        payments["due_dt"] = dt(payments.get("due_date", ""))
        payments["overdue_days"] = (pd.Timestamp.today().normalize() - payments["due_dt"]).dt.days.fillna(0).clip(lower=0)
    else:
        pay_sum = pd.DataFrame(columns=["project_name","received","contract"])
        payments["overdue_days"] = []

    p = projects.merge(labor, on="project_name", how="left").merge(travel_sum, on="project_name", how="left").merge(pay_sum, on="project_name", how="left")
    for c in ["labor_cost","travel_cost","received","contract"]:
        p[c] = num(p.get(c,0))
    p["total_cost"] = p["labor_cost"] + p["travel_cost"]
    p["gross_profit"] = p["revenue"] - p["total_cost"]
    p["gross_margin"] = p.apply(lambda r: r["gross_profit"] / r["revenue"] if r["revenue"] else 0, axis=1)
    p["unreceived"] = p["revenue"] - p["received"]
    p["recovery_rate"] = p.apply(lambda r: r["received"] / r["revenue"] if r["revenue"] else 0, axis=1)
    p["risk_level"] = p["gross_margin"].apply(lambda x: "亏损" if x < 0 else ("红色" if x < .1 else ("橙色" if x < .25 else ("黄色" if x < .4 else "绿色"))))

    if not capacity.empty:
        capacity = capacity.merge(people[["name","standard_monthly_visits"]], left_on="auditor", right_on="name", how="left") if not people.empty else capacity
        capacity["standard"] = num(capacity.get("standard_monthly_visits",0))
        capacity.loc[capacity["standard"].eq(0), "standard"] = 5
        capacity["utilization"] = capacity["visits"] / capacity["standard"]
        capacity["capacity_status"] = capacity["utilization"].apply(lambda x: "超负荷" if x > 1.1 else ("饱和" if x >= .9 else ("正常" if x >= .6 else "产能不足")))

    customers = p.groupby("customer", as_index=False).agg(revenue=("revenue","sum"), cost=("total_cost","sum"), profit=("gross_profit","sum"), received=("received","sum"), project_count=("project_name","count")) if not p.empty else pd.DataFrame(columns=["customer","revenue","cost","profit","received","project_count"])
    if not customers.empty:
        customers["margin"] = customers.apply(lambda r: r.profit/r.revenue if r.revenue else 0, axis=1)
        customers["recovery_rate"] = customers.apply(lambda r: r.received/r.revenue if r.revenue else 0, axis=1)
        customers["grade"] = customers.apply(lambda r: "A类高价值" if r.margin>=.4 and r.recovery_rate>=.8 else ("B类优化报价" if r.margin>=.25 else "D类风险客户"), axis=1)

    return p, payments, customers, capacity


def money(x): return f"¥{float(x or 0):,.0f}"
def pct(x): return f"{float(x or 0)*100:.1f}%"


def kpi(label, val, help_text=""):
    st.metric(label, val, help=help_text)


def login_page():
    st.title(APP_TITLE)
    st.caption("Enterprise V2.0 | 经营、利润、现金流、报价、产能一体化")
    with st.form("login"):
        u = st.text_input("用户名", value="admin")
        p = st.text_input("密码", type="password")
        ok = st.form_submit_button("登录", use_container_width=True)
    if ok:
        users = q("SELECT * FROM users WHERE username=? AND active=1", (u,))
        if not users.empty and verify_pw(p, users.iloc[0]["password_hash"]):
            st.session_state.user = u
            st.session_state.role = users.iloc[0]["role"]
            log("登录", "成功")
            st.rerun()
        else:
            st.error("用户名或密码错误")


def sidebar():
    st.sidebar.title(APP_TITLE)
    st.sidebar.caption(f"当前用户：{st.session_state.user}｜{st.session_state.role}")
    pages = ["老板驾驶舱","项目经营分析","现金流回款","报价测算","客户利润贡献","人员产能分析","预警中心","数据导入维护","老板月报","系统设置"]
    page = st.sidebar.radio("导航", pages)
    if st.sidebar.button("退出登录"):
        st.session_state.clear(); st.rerun()
    return page


def page_dashboard():
    st.header("老板经营驾驶舱")
    p, payments, customers, capacity = calc_all()
    cols = st.columns(4)
    cols[0].metric("确认收入", money(p["revenue"].sum() if not p.empty else 0))
    cols[1].metric("已回款", money(p["received"].sum() if not p.empty else 0))
    cols[2].metric("未回款", money(p["unreceived"].sum() if not p.empty else 0))
    margin = (p["gross_profit"].sum()/p["revenue"].sum()) if not p.empty and p["revenue"].sum() else 0
    cols[3].metric("综合毛利率", pct(margin))
    cols = st.columns(4)
    cols[0].metric("项目数", len(p))
    cols[1].metric("项目毛利", money(p["gross_profit"].sum() if not p.empty else 0))
    cols[2].metric("风险项目", int((p["risk_level"].isin(["亏损","红色","橙色"])).sum()) if not p.empty else 0)
    cols[3].metric("超负荷人员", int((capacity.get("capacity_status",pd.Series(dtype=str))=="超负荷").sum()) if not capacity.empty else 0)
    st.subheader("项目毛利排行")
    st.dataframe(p[["project_name","customer","revenue","total_cost","gross_profit","gross_margin","risk_level"]].sort_values("gross_profit", ascending=False), use_container_width=True)
    if not p.empty:
        st.bar_chart(p.set_index("project_name")[["revenue","total_cost","gross_profit"]])


def page_projects():
    st.header("项目经营分析")
    p, *_ = calc_all()
    st.dataframe(p, use_container_width=True)


def page_cashflow():
    st.header("现金流与回款周期")
    p, payments, *_ = calc_all()
    if not payments.empty:
        payments["due_dt"] = dt(payments.get("due_date",""))
        payments["received_dt"] = dt(payments.get("received_date",""))
        payments["cycle_days"] = (payments["received_dt"] - payments["due_dt"]).dt.days
        st.dataframe(payments, use_container_width=True)
    st.subheader("项目回款汇总")
    st.dataframe(p[["project_name","customer","revenue","received","unreceived","recovery_rate"]], use_container_width=True)


def page_quote():
    st.header("报价测算模型")
    c1,c2,c3,c4 = st.columns(4)
    people_n = c1.number_input("每院次人数", 1, 10, 2)
    days_n = c2.number_input("每院次天数", .5, 10.0, 2.0)
    visits = c3.number_input("院次数", 1, 100, 1)
    daily_cost = c4.number_input("单人日成本", 0.0, 10000.0, 800.0)
    travel_per_visit = st.number_input("单院次差旅预算", 0.0, 100000.0, 6000.0)
    tax_rate = st.number_input("税率", 0.0, .3, .06)
    target_margin = st.number_input("目标毛利率", 0.0, .9, .4)
    risk_buffer = st.number_input("风险缓冲率", 0.0, .5, .08)
    labor = people_n * days_n * visits * daily_cost
    travel = travel_per_visit * visits
    cost = (labor + travel) * (1 + risk_buffer)
    price = cost / (1 - target_margin) if target_margin < 1 else 0
    price_with_tax = price * (1 + tax_rate)
    st.success(f"建议报价：{money(price_with_tax)}；成本底线：{money(cost*(1+tax_rate))}；预计总成本：{money(cost)}")
    st.info("两人两天含税含差旅项目建议设置公司报价红线，低于成本底线必须审批。")


def page_customers():
    st.header("客户利润贡献")
    _,_,customers,_ = calc_all()
    st.dataframe(customers.sort_values("profit", ascending=False) if not customers.empty else customers, use_container_width=True)


def page_capacity():
    st.header("人员产能利用率")
    *_, capacity = calc_all()
    st.dataframe(capacity, use_container_width=True)


def page_alerts():
    st.header("项目盈亏预警中心")
    p, payments, _, capacity = calc_all()
    alerts = []
    for _, r in p.iterrows():
        if r.get("gross_profit",0) < 0: alerts.append([r.project_name, r.customer, "亏损预警", money(r.gross_profit), "红色", "复盘报价、成本与差旅"])
        elif r.get("gross_margin",0) < .25: alerts.append([r.project_name, r.customer, "低毛利预警", pct(r.gross_margin), "橙色", "后续报价上调或控制交付成本"])
        if r.get("recovery_rate",0) < .5 and r.get("revenue",0) > 0: alerts.append([r.project_name, r.customer, "回款不足", pct(r.recovery_rate), "黄色", "商务/财务跟进回款"])
    for _, r in capacity.iterrows() if not capacity.empty else []:
        if r.get("capacity_status") == "超负荷": alerts.append([r.auditor, "人员", "产能过载", pct(r.utilization), "红色", "调整排班或增加兼职支持"])
    st.dataframe(pd.DataFrame(alerts, columns=["对象","客户/类型","风险类型","当前情况","风险等级","建议动作"]), use_container_width=True)


def page_import():
    st.header("数据导入维护")
    table = st.selectbox("选择数据表", list(TABLES.keys()), format_func=lambda x: TABLES[x]["label"])
    st.caption("支持 CSV / Excel。系统会自动识别常见中文列名，并映射到标准字段。")
    up = st.file_uploader("上传文件", type=["csv","xlsx","xls"])
    if up:
        raw = load_upload(up)
        if table == "travel" and st.checkbox("启用财务差旅自动识别", value=True):
            df = auto_recognize_travel(raw)
        else:
            df = normalize_columns(raw, table)
        edited = st.data_editor(df, use_container_width=True, num_rows="dynamic")
        mode = st.radio("写入方式", ["追加","覆盖"], horizontal=True)
        if st.button("写入数据库", type="primary") and has_perm("import"):
            write_table(table, pd.DataFrame(edited), "replace" if mode=="覆盖" else "append")
            st.success("已写入数据库")
    st.subheader("当前数据")
    current = read_table(table)
    edited2 = st.data_editor(current, use_container_width=True, num_rows="dynamic", key=f"edit_{table}")
    if st.button("保存当前表修改"):
        if has_perm("edit"):
            write_table(table, pd.DataFrame(edited2), "replace"); st.success("已保存")
        else: st.error("当前账号无编辑权限")
    if has_perm("export"):
        buf = io.BytesIO(); current.to_excel(buf, index=False)
        st.download_button("导出当前表 Excel", buf.getvalue(), f"{table}.xlsx")


def build_report() -> bytes:
    p, payments, customers, capacity = calc_all()
    doc = Document(); doc.add_heading("万宁睿和经营分析月报", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    rev = p["revenue"].sum() if not p.empty else 0; cost = p["total_cost"].sum() if not p.empty else 0; profit = p["gross_profit"].sum() if not p.empty else 0
    doc.add_heading("一、经营概览", 1)
    for line in [f"确认收入：{money(rev)}", f"总成本：{money(cost)}", f"项目毛利：{money(profit)}", f"综合毛利率：{pct(profit/rev if rev else 0)}", f"已回款：{money(p['received'].sum() if not p.empty else 0)}", f"未回款：{money(p['unreceived'].sum() if not p.empty else 0)}"]: doc.add_paragraph(line)
    doc.add_heading("二、关键判断", 1)
    risk_n = int((p.get("risk_level",pd.Series(dtype=str)).isin(["亏损","红色","橙色"])).sum()) if not p.empty else 0
    doc.add_paragraph(f"本期风险项目 {risk_n} 个。建议管理层重点关注低毛利、回款不足和人员超负荷项目。")
    for title, df, cols in [("三、项目经营明细",p,["project_name","customer","revenue","total_cost","gross_profit","gross_margin","risk_level"]),("四、客户利润贡献",customers,["customer","revenue","cost","profit","margin","recovery_rate","grade"]),("五、人员产能",capacity,["auditor","visits","travel_days","utilization","capacity_status"] )]:
        doc.add_heading(title, 1)
        if df.empty: doc.add_paragraph("暂无数据"); continue
        t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
        for i,c in enumerate(cols): t.rows[0].cells[i].text = c
        for _, r in df[cols].head(30).iterrows():
            cells = t.add_row().cells
            for i,c in enumerate(cols): cells[i].text = str(r.get(c,""))
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()


def page_report():
    st.header("老板月报自动生成")
    st.write("自动汇总经营概览、项目毛利、客户贡献、人员产能与管理建议，生成 Word 月报。")
    if st.button("生成老板月报", type="primary"):
        data = build_report()
        st.download_button("下载经营分析月报 Word", data, f"经营分析月报_{date.today()}.docx")


def page_settings():
    st.header("系统设置")
    tabs = st.tabs(["用户与权限","操作日志","数据导出"])
    with tabs[0]:
        users = q("SELECT username, role, active FROM users")
        st.dataframe(users, use_container_width=True)
        if has_perm("all"):
            with st.form("add_user"):
                u = st.text_input("用户名")
                p = st.text_input("密码", type="password")
                r = st.selectbox("角色", list(ROLES.keys()))
                if st.form_submit_button("新增/重置用户"):
                    exec_sql("REPLACE INTO users(username,password_hash,role,active) VALUES(?,?,?,1)", (u, hash_pw(p), r)); st.success("已保存")
    with tabs[1]:
        st.dataframe(q("SELECT * FROM audit_logs ORDER BY ts DESC LIMIT 200"), use_container_width=True)
    with tabs[2]:
        if has_perm("export") and st.button("导出全部数据 ZIP"):
            zbuf = io.BytesIO()
            with zipfile.ZipFile(zbuf, "w") as z:
                for t in TABLES:
                    bio = io.BytesIO(); read_table(t).to_excel(bio, index=False); z.writestr(f"{t}.xlsx", bio.getvalue())
            st.download_button("下载ZIP", zbuf.getvalue(), "business_dashboard_data.zip")


def main():
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    init_db()
    if "user" not in st.session_state:
        login_page(); return
    page = sidebar()
    if page != "老板驾驶舱" and not has_perm("view"):
        st.error("当前账号无查看权限"); return
    {"老板驾驶舱":page_dashboard,"项目经营分析":page_projects,"现金流回款":page_cashflow,"报价测算":page_quote,"客户利润贡献":page_customers,"人员产能分析":page_capacity,"预警中心":page_alerts,"数据导入维护":page_import,"老板月报":page_report,"系统设置":page_settings}[page]()

if __name__ == "__main__":
    main()
